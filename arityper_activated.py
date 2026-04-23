#!/usr/bin/env python3
"""
AriTyper Activated — Premium UI Redesign
Activation-first flow with glassmorphism dark theme, smooth animations,
professional typography and rounded card layout.
All backend logic preserved from original.
"""
import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog, messagebox
import requests
import threading
import time
import json
import os
import sys
import socket
import hashlib
import platform
from datetime import datetime

# Windows-specific imports for enhanced functionality
try:
    import pywintypes
    import win32gui
    import win32con
    import win32api
    import win32clipboard
    import win32com.client
    import pythoncom
    from PIL import Image, ImageTk
    import io
    WINDOWS_SUPPORT = True
except ImportError:
    WINDOWS_SUPPORT = False
    # Create dummy objects if imports fail
    pywintypes = None
    win32gui = None
    win32con = None
    win32api = None
    win32clipboard = None
    win32com = None
    pythoncom = None

# Document processing imports (always available)
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# ═══════════════════════════════════════════════════════
#  DESIGN TOKENS
# ═══════════════════════════════════════════════════════
BG       = "#060b14"
SURFACE  = "#0b1623"
CARD     = "#0f1f31"
CARD2    = "#132538"
BORDER   = "#1a3050"
ACCENT   = "#00d4ff"
ACCENT2  = "#0066cc"
GREEN    = "#00e676"
ORANGE   = "#ffab40"
RED      = "#ff5252"
YELLOW   = "#ffd740"
TEXT     = "#ddeeff"
MUTED    = "#3d6080"
DIM      = "#1e3a55"

# Typography — Palatino for display, Verdana for body
FT_HERO  = ("Palatino", 28, "bold")
FT_HEAD  = ("Palatino", 17, "bold")
FT_SUB   = ("Verdana", 11, "bold")
FT_BODY  = ("Verdana", 9)
FT_MONO  = ("Courier New", 11, "bold")
FT_BTN   = ("Verdana", 11, "bold")
FT_SM    = ("Verdana", 8)
FT_CAPS  = ("Verdana", 8, "bold")


# ═══════════════════════════════════════════════════════
#  REUSABLE WIDGETS
# ═══════════════════════════════════════════════════════

class RoundedCard(tk.Frame):
    """Dark card with 1-px accent border achieved via outer/inner frame."""
    def __init__(self, parent, border_color=BORDER, **kw):
        outer = tk.Frame(parent, bg=border_color, bd=0)
        super().__init__(outer, bg=CARD, bd=0, **kw)
        self.pack(fill="both", expand=True, padx=1, pady=1)
        self._outer = outer

    def place_in(self, **pack_kw):
        self._outer.pack(**pack_kw)
        return self


class PulseLabel(tk.Label):
    """Label that alternates between two colours."""
    def __init__(self, *a, c1=ORANGE, c2=MUTED, interval=750, **kw):
        super().__init__(*a, **kw)
        self._c = [c1, c2]
        self._i = 0
        self._interval = interval
        self._active = False

    def start(self):
        self._active = True
        self._tick()

    def stop(self, color=None):
        self._active = False
        if color:
            self.config(fg=color)

    def _tick(self):
        if not self._active:
            return
        self.config(fg=self._c[self._i % 2])
        self._i += 1
        self.after(self._interval, self._tick)


class GlowButton(tk.Button):
    """Button with hover highlight and active-press effect."""
    def __init__(self, *a, base_bg=ACCENT, hover_bg=None, **kw):
        self._base  = base_bg
        self._hover = hover_bg or self._lighten(base_bg)
        super().__init__(*a, bg=base_bg, activebackground=self._hover,
                         relief="flat", bd=0, **kw)
        self.bind("<Enter>", self._on_enter)
        self.bind("<Leave>", self._on_leave)

    def _on_enter(self, _):
        if str(self["state"]) != "disabled":
            self.config(bg=self._hover)

    def _on_leave(self, _):
        if str(self["state"]) != "disabled":
            self.config(bg=self._base)

    @staticmethod
    def _lighten(hex_color):
        r = min(255, int(hex_color[1:3], 16) + 24)
        g = min(255, int(hex_color[3:5], 16) + 24)
        b = min(255, int(hex_color[5:7], 16) + 24)
        return f"#{r:02x}{g:02x}{b:02x}"


class ProgressArc(tk.Canvas):
    """Thin rounded progress bar drawn on a Canvas."""
    def __init__(self, parent, w=460, h=10, **kw):
        super().__init__(parent, width=w, height=h,
                         bg=CARD, highlightthickness=0, **kw)
        self._w, self._h = w, h
        self._pct = 0
        self._draw()

    def set_pct(self, v):
        self._pct = max(0, min(100, v))
        self._draw()

    def _draw(self):
        self.delete("all")
        r = self._h // 2
        # track
        self._rr(0, 0, self._w, self._h, r, fill=CARD2, outline=DIM)
        # fill
        fw = int(self._w * self._pct / 100)
        if fw >= self._h:
            self._rr(0, 0, fw, self._h, r, fill=ACCENT, outline="")
        elif fw > 0:
            self.create_oval(0, 0, self._h, self._h, fill=ACCENT, outline="")

    def _rr(self, x1, y1, x2, y2, r, **kw):
        pts = [x1+r,y1, x2-r,y1, x2,y1, x2,y1+r,
               x2,y2-r, x2,y2, x2-r,y2, x1+r,y2,
               x1,y2, x1,y2-r, x1,y1+r, x1,y1]
        self.create_polygon(pts, smooth=True, **kw)


def _sep(parent, bg=BORDER, pady=6):
    tk.Frame(parent, height=1, bg=bg).pack(fill="x", padx=20, pady=pady)


def _section(parent, icon, title, bg=CARD):
    row = tk.Frame(parent, bg=bg)
    row.pack(fill="x", padx=20, pady=(14, 4))
    tk.Label(row, text=icon, font=("Segoe UI Emoji", 13),
             bg=bg, fg=ACCENT).pack(side="left")
    tk.Label(row, text=f"  {title}", font=FT_SUB,
             bg=bg, fg=TEXT).pack(side="left")
    _sep(parent, pady=4)


# ═══════════════════════════════════════════════════════
#  MAIN APP CLASS
# ═══════════════════════════════════════════════════════

class AriTyperActivated:
    """AriTyper with activation-first UI — premium redesign."""

    def __init__(self, root):
        self.root = root
        self.server_url = "http://localhost:5000"
        self.device_id = self._generate_device_id()
        self.license_data = None
        self.target_window = None
        self.is_typing = False
        self.current_view = "activation"

        self._setup_window()
        self._style_ttk()
        self.create_activation_ui()
        self.start_app()

    # ───────────────────────────────────────────
    #  WINDOW
    # ───────────────────────────────────────────
    def _setup_window(self):
        self.root.title("AriTyper  ·  Activation")
        self.root.geometry("920x720")
        self.root.configure(bg=BG)
        self.root.resizable(True, True)
        self.root.minsize(720, 540)
        self.root.update_idletasks()
        x = (self.root.winfo_screenwidth()  // 2) - 460
        y = (self.root.winfo_screenheight() // 2) - 360
        self.root.geometry(f"920x720+{x}+{y}")

    def _style_ttk(self):
        s = ttk.Style()
        s.theme_use("clam")
        s.configure("TScrollbar",
                    troughcolor=SURFACE, background=BORDER,
                    arrowcolor=MUTED,   bordercolor=SURFACE)
        s.configure("Horizontal.TScale",
                    background=CARD2, troughcolor=DIM, slidercolor=ACCENT)

    # ───────────────────────────────────────────
    #  SHARED HEADER
    # ───────────────────────────────────────────
    def _build_header(self, parent, tag="Activation Required"):
        hdr = tk.Frame(parent, bg=SURFACE, height=82)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)

        # left colour bar
        tk.Frame(hdr, width=5, bg=ACCENT).pack(side="left", fill="y")

        left = tk.Frame(hdr, bg=SURFACE)
        left.pack(side="left", fill="both", expand=True, padx=20)

        tk.Label(left, text="⌨  AriTyper",
                 font=FT_HERO, bg=SURFACE, fg=ACCENT).pack(anchor="w", pady=(10, 0))
        tk.Label(left, text="Professional Document Typing Software",
                 font=FT_SM, bg=SURFACE, fg=MUTED).pack(anchor="w")

        # right tag badge
        badge = tk.Frame(hdr, bg=DIM, padx=12, pady=4)
        badge.pack(side="right", padx=20, pady=22)
        tk.Label(badge, text=tag, font=FT_CAPS, bg=DIM, fg=ACCENT).pack()

    # ═══════════════════════════════════════════
    #  ACTIVATION UI
    # ═══════════════════════════════════════════
    def create_activation_ui(self):
        self.activation_frame = tk.Frame(self.root, bg=BG)
        self.activation_frame.pack(fill="both", expand=True)

        self._build_header(self.activation_frame)

        # ── scrollable body ──────────────────────
        wrap = tk.Frame(self.activation_frame, bg=BG)
        wrap.pack(fill="both", expand=True)

        canvas = tk.Canvas(wrap, bg=BG, highlightthickness=0)
        vsb = ttk.Scrollbar(wrap, orient="vertical", command=canvas.yview)
        sf  = tk.Frame(canvas, bg=BG)
        sf.bind("<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=sf, anchor="nw")
        canvas.configure(yscrollcommand=vsb.set)
        canvas.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        canvas.bind("<MouseWheel>",
            lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        # ── cards with better spacing ────────────────────────
        self._card_device(sf)
        self._card_license_status(sf)
        self._card_payment(sf)
        self._card_transaction(sf)

        self.status_label = tk.Label(
            sf, text="Connecting to server…",
            font=FT_SM, bg=BG, fg=MUTED)
        self.status_label.pack(pady=(6, 24))

    # ── Device card ─────────────────────────────
    def _card_device(self, parent):
        c = RoundedCard(parent)
        c.place_in(fill="x", padx=32, pady=(20, 8))

        _section(c, "🖥", "Device Information")
        row = tk.Frame(c, bg=CARD)
        row.pack(fill="x", padx=20, pady=(0, 14))
        tk.Label(row, text="Device ID:", font=FT_BODY,
                 bg=CARD, fg=MUTED).pack(side="left")
        tk.Label(row, text=f"  {self.device_id}",
                 font=FT_MONO, bg=CARD, fg=ACCENT).pack(side="left")

    # ── License status card ──────────────────────
    def _card_license_status(self, parent):
        c = RoundedCard(parent)
        c.place_in(fill="x", padx=32, pady=8)

        self.license_status_label = PulseLabel(
            c, text="🔒  License Status: Checking…",
            font=FT_SUB, bg=CARD, fg=ORANGE,
            c1=ORANGE, c2=MUTED)
        self.license_status_label.pack(pady=18)
        self.license_status_label.start()

    # ── Payment card ─────────────────────────────
    def _card_payment(self, parent):
        c = RoundedCard(parent)
        c.place_in(fill="x", padx=32, pady=8)

        _section(c, "💳", "Purchase License — UGX 10,000")

        # Tab-style selector
        tab_row = tk.Frame(c, bg=CARD)
        tab_row.pack(fill="x", padx=20, pady=(0, 10))

        self._active_tab   = tk.StringVar(value="mtn")
        self._mtn_content  = None
        self._air_content  = None

        self._tab_mtn = tk.Button(
            tab_row, text="📱  MTN MoMo",
            font=FT_BTN, bg=YELLOW, fg=BG,
            relief="flat", bd=0, padx=16, pady=8,
            cursor="hand2",
            command=lambda: self._switch_tab("mtn"))
        self._tab_mtn.pack(side="left", padx=(0, 6))

        self._tab_air = tk.Button(
            tab_row, text="📱  Airtel Money",
            font=FT_BTN, bg=CARD2, fg=MUTED,
            relief="flat", bd=0, padx=16, pady=8,
            cursor="hand2",
            command=lambda: self._switch_tab("airtel"))
        self._tab_air.pack(side="left")

        # Tab content container
        self._tab_body = tk.Frame(c, bg=CARD)
        self._tab_body.pack(fill="x", padx=20, pady=(0, 8))

        self._build_mtn_tab()
        self._build_airtel_tab()
        self._switch_tab("mtn")

        # After payment
        _sep(c, pady=6)
        self._after_payment_block(c)

    def _build_mtn_tab(self):
        f = tk.Frame(self._tab_body, bg=CARD2, padx=16, pady=12)
        steps = [
            ("1", "Dial  *165#  on your MTN line",              TEXT),
            ("2", "Select  'Pay a Bill'  or  'Send Money'",     MUTED),
            ("3", "Merchant ID:   7074948",                     ACCENT),
            ("4", "Amount:   10,000 UGX",                       GREEN),
            ("5", "Confirm with PIN — note your Transaction ID", MUTED),
        ]
        self._render_steps(f, steps)
        self._mtn_content = f

    def _build_airtel_tab(self):
        f = tk.Frame(self._tab_body, bg=CARD2, padx=16, pady=12)
        steps = [
            ("1", "Dial  *185#  on your Airtel line",            TEXT),
            ("2", "Select  'Make Payments'  or  'Send Money'",   MUTED),
            ("3", "Merchant ID:   66562536",                     ACCENT),
            ("4", "Amount:   10,000 UGX",                        GREEN),
            ("5", "Confirm with PIN — note your Transaction ID",  MUTED),
        ]
        self._render_steps(f, steps)
        self._air_content = f

    def _render_steps(self, parent, steps):
        for num, txt, col in steps:
            row = tk.Frame(parent, bg=CARD2)
            row.pack(fill="x", pady=2)
            tk.Label(row, text=f"  {num}.", font=FT_SM,
                     bg=CARD2, fg=MUTED, width=3, anchor="e").pack(side="left")
            tk.Label(row, text=f"  {txt}", font=FT_BODY,
                     bg=CARD2, fg=col, anchor="w").pack(side="left")

    def _switch_tab(self, tab):
        self._active_tab.set(tab)
        for w in self._tab_body.winfo_children():
            w.pack_forget()

        if tab == "mtn":
            self._mtn_content.pack(fill="x")
            self._tab_mtn.config(bg=YELLOW,  fg=BG)
            self._tab_air.config(bg=CARD2,   fg=MUTED)
        else:
            self._air_content.pack(fill="x")
            self._tab_air.config(bg=ORANGE,  fg=BG)
            self._tab_mtn.config(bg=CARD2,   fg=MUTED)

    def _after_payment_block(self, parent):
        bx = tk.Frame(parent, bg=CARD, padx=20, pady=10)
        bx.pack(fill="x")
        tk.Label(bx, text="✅  After Payment",
                 font=FT_SUB, bg=CARD, fg=GREEN).pack(anchor="w", pady=(0, 6))
        for line in [
            "Business:  AT Tech Solutions Uganda",
            "WhatsApp:  +256 760 730 254  (send Transaction ID + your name)",
            "Approval usually takes just a few minutes.",
            "The app will unlock automatically once approved.",
        ]:
            tk.Label(bx, text=line, font=FT_BODY,
                     bg=CARD, fg=MUTED, anchor="w").pack(fill="x")

    # ── Transaction card ─────────────────────────
    def _card_transaction(self, parent):
        c = RoundedCard(parent, border_color=ACCENT)
        c.place_in(fill="x", padx=32, pady=(8, 20))

        _section(c, "🔑", "Enter Your Transaction ID")

        tk.Label(c,
                 text="After payment, paste your transaction ID below for instant approval.",
                 font=FT_BODY, bg=CARD, fg=MUTED,
                 wraplength=760, justify="left").pack(
                     anchor="w", padx=20, pady=(0, 10))

        # Entry with glow border
        ef = tk.Frame(c, bg=ACCENT, padx=2, pady=2)
        ef.pack(fill="x", padx=20, pady=(0, 12))
        self.transaction_entry = tk.Entry(
            ef, font=("Courier New", 13, "bold"),
            bg=CARD2, fg=WHITE if False else "#ccddee",
            insertbackground=ACCENT, relief="flat", bd=0)
        self.transaction_entry.pack(fill="x", ipady=10)
        self.transaction_entry.insert(0, "Paste transaction ID here…")
        self.transaction_entry.config(fg=MUTED)
        self.transaction_entry.bind("<FocusIn>",  self._entry_in)
        self.transaction_entry.bind("<FocusOut>", self._entry_out)

        # Submit button
        self.submit_btn = GlowButton(
            c,
            text="📤   Submit Transaction ID",
            font=FT_BTN,
            base_bg=ACCENT, fg=BG,
            padx=24, pady=13,
            cursor="hand2",
            command=self.submit_transaction)
        self.submit_btn.pack(fill="x", padx=20, pady=(0, 6))

        self.txn_status = tk.Label(
            c, text="", font=FT_SM, bg=CARD, fg=MUTED)
        self.txn_status.pack(pady=(0, 14))

    def _entry_in(self, _):
        if "Paste" in self.transaction_entry.get():
            self.transaction_entry.delete(0, "end")
            self.transaction_entry.config(fg=TEXT)

    def _entry_out(self, _):
        if not self.transaction_entry.get().strip():
            self.transaction_entry.insert(0, "Paste transaction ID here…")
            self.transaction_entry.config(fg=MUTED)

    # ═══════════════════════════════════════════
    #  MAIN UI (post-activation)
    # ═══════════════════════════════════════════
    def create_main_ui(self):
        self.activation_frame.pack_forget()
        self.root.title("AriTyper  ·  Licensed")

        self.main_frame = tk.Frame(self.root, bg=BG)
        self.main_frame.pack(fill="both", expand=True)

        self._build_header(self.main_frame, tag="✅  Licensed")

        # active pill bar
        pill = tk.Frame(self.main_frame, bg=SURFACE, height=30)
        pill.pack(fill="x")
        pill.pack_propagate(False)
        tk.Label(pill, text=f"● ACTIVE  |  {self.device_id[:22]}…",
                 font=FT_CAPS, bg=SURFACE, fg=GREEN, padx=16).pack(side="left", pady=6)

        # ── grid layout body ──────────────────────
        body = tk.Frame(self.main_frame, bg=BG)
        body.pack(fill="both", expand=True, padx=28, pady=18)

        # Configure grid columns for better space usage
        body.grid_columnconfigure(0, weight=1)
        body.grid_columnconfigure(1, weight=1)
        body.grid_rowconfigure(0, weight=1)
        body.grid_rowconfigure(1, weight=1)

        # Create panels with grid
        left = tk.Frame(body, bg=BG)
        left.grid(row=0, column=0, sticky="nsew", padx=(0, 8), pady=(0, 4))
        
        right = tk.Frame(body, bg=BG)
        right.grid(row=0, column=1, sticky="nsew", padx=(8, 0), pady=(0, 4))

        # Document display panel (spans full width)
        doc_panel = tk.Frame(body, bg=BG)
        doc_panel.grid(row=1, column=0, columnspan=2, sticky="nsew", pady=(4, 0))

        self._build_file_panel(left)
        self._build_controls_panel(right)
        self._build_document_panel(doc_panel)

        self.main_status_label = tk.Label(
            self.main_frame, text="Ready.",
            font=FT_BODY, bg=BG, fg=MUTED)
        self.main_status_label.pack(pady=(0, 10))

    # ── File panel ───────────────────────────────
    def _build_file_panel(self, parent):
        c = RoundedCard(parent)
        c._outer.pack(fill="both", expand=True)

        _section(c, "📄", "Document & Content Input")

        # Quick actions row - prominent paste and window selection
        quick_actions = tk.Frame(c, bg=CARD)
        quick_actions.pack(fill="x", padx=20, pady=(0, 12))

        # Window selection button (prominent)
        self.window_btn = GlowButton(
            quick_actions, text="🪟 Select Window",
            font=FT_BTN, base_bg=ACCENT2, fg=TEXT,
            padx=16, pady=10, cursor="hand2",
            command=self.select_target_window
        )
        self.window_btn.pack(side="left", padx=(0, 8))

        # Paste button (prominent)
        self.paste_btn = GlowButton(
            quick_actions, text="📋 Paste Text",
            font=FT_BTN, base_bg=ORANGE, fg=TEXT,
            padx=16, pady=10, cursor="hand2",
            command=self.paste_text
        )
        self.paste_btn.pack(side="left", padx=(0, 8))

        # Clear button
        self.clear_btn = GlowButton(
            quick_actions, text="🗑️ Clear",
            font=FT_BTN, base_bg=RED, fg=TEXT,
            padx=16, pady=10, cursor="hand2",
            command=self.clear_document
        )
        self.clear_btn.pack(side="right")

        # Window info label
        self.window_info = tk.Label(
            c, text="No target window selected",
            font=FT_SM, bg=CARD, fg=MUTED
        )
        self.window_info.pack(anchor="w", padx=20, pady=(0, 8))

        _sep(c, pady=8)

        # drop zone
        dz = tk.Frame(c, bg=CARD2, cursor="hand2", padx=10, pady=10)
        dz.pack(fill="x", padx=20, pady=(0, 12))
        dz.bind("<Button-1>", lambda e: self.browse_file())

        tk.Label(dz, text="📂", font=("Segoe UI Emoji", 30),
                 bg=CARD2, fg=MUTED).pack(pady=(16, 4))
        self.file_label = tk.Label(
            dz, text="Click to select or drag a file",
            font=FT_BODY, bg=CARD2, fg=MUTED, wraplength=280)
        self.file_label.pack(pady=(0, 4))
        tk.Label(dz, text="TXT  ·  DOCX  ·  PDF",
                 font=FT_SM, bg=CARD2, fg=DIM).pack(pady=(0, 18))

        GlowButton(
            c, text="Browse Files",
            font=FT_BTN, base_bg=ACCENT2, fg=TEXT,
            padx=20, pady=10, cursor="hand2",
            command=self.browse_file
        ).pack(padx=20, pady=(0, 18))

    # ── Controls panel ───────────────────────────
    def _build_controls_panel(self, parent):
        c = RoundedCard(parent)
        c._outer.pack(fill="both", expand=True)

        _section(c, "⌨", "Typing Controls")

        # Speed
        sr = tk.Frame(c, bg=CARD)
        sr.pack(fill="x", padx=20, pady=(0, 8))
        tk.Label(sr, text="Speed", font=FT_BODY, bg=CARD, fg=MUTED).pack(side="left")
        self.speed_var = tk.DoubleVar(value=50)
        ttk.Scale(sr, from_=1, to=100, variable=self.speed_var,
                  orient="horizontal", length=170,
                  command=self._update_speed_label
                  ).pack(side="left", padx=10)
        self.speed_label = tk.Label(sr, text="50 %",
                                    font=FT_MONO, bg=CARD, fg=ACCENT)
        self.speed_label.pack(side="left")

        # Progress bar
        pb_wrap = tk.Frame(c, bg=CARD)
        pb_wrap.pack(fill="x", padx=20, pady=(10, 4))
        tk.Label(pb_wrap, text="Progress", font=FT_SM,
                 bg=CARD, fg=MUTED).pack(anchor="w", pady=(0, 4))
        self.progress_bar = ProgressArc(pb_wrap, w=360, h=12)
        self.progress_bar.pack()

        self._pct_lbl = tk.Label(c, text="0 %", font=FT_SM, bg=CARD, fg=MUTED)
        self._pct_lbl.pack(anchor="e", padx=24, pady=(2, 10))

        _sep(c, pady=6)

        # Buttons
        br = tk.Frame(c, bg=CARD)
        br.pack(padx=20, pady=(6, 18))

        self.start_btn = GlowButton(
            br, text="▶  Start Typing",
            font=FT_BTN, base_bg=GREEN, fg=BG,
            padx=26, pady=13, cursor="hand2",
            command=self.start_typing)
        self.start_btn.pack(side="left", padx=(0, 8))

        self.stop_btn = GlowButton(
            br, text="⏹  Stop",
            font=FT_BTN, base_bg=RED, fg=TEXT,
            padx=26, pady=13, cursor="hand2",
            state="disabled", command=self.stop_typing)
        self.stop_btn.pack(side="left")

    # ── Document panel ───────────────────────────
    def _build_document_panel(self, parent):
        c = RoundedCard(parent)
        c._outer.pack(fill="both", expand=True)

        _section(c, "📖", "Document Content Preview")

        # Document display area
        self.doc_frame = tk.Frame(c, bg=CARD2)
        self.doc_frame.pack(fill="both", expand=True, padx=20, pady=(0, 12))

        # Document text display
        self.doc_text = scrolledtext.ScrolledText(
            self.doc_frame, 
            font=("Consolas", 11),
            bg=CARD2, fg=TEXT,
            wrap=tk.WORD,
            height=8,
            relief="flat",
            bd=0,
            padx=12, pady=8
        )
        self.doc_text.pack(fill="both", expand=True)
        self.doc_text.config(state="disabled")

        # Status info
        status_frame = tk.Frame(c, bg=CARD)
        status_frame.pack(fill="x", padx=20, pady=(0, 18))
        
        self.doc_status = tk.Label(
            status_frame, text="No content loaded - paste text or select a file",
            font=FT_SM, bg=CARD, fg=MUTED
        )
        self.doc_status.pack(anchor="w")

    # ═══════════════════════════════════════════
    #  BACKEND  (original logic — unchanged)
    # ═══════════════════════════════════════════
    def start_app(self):
        threading.Thread(target=self._start_device_monitoring, daemon=True).start()
        threading.Thread(target=self._start_license_monitoring, daemon=True).start()
        self._check_license()

    def _start_device_monitoring(self):
        time.sleep(1)
        self._register_device()
        self._start_heartbeat()

    def _register_device(self):
        try:
            payload = {
                "device_id": self.device_id,
                "hostname":  platform.node(),
                "os_info":   f"{platform.system()} {platform.release()}",
                "user_info": {"app_version": "2.0.0"}
            }
            r = requests.post(f"{self.server_url}/api/device_heartbeat",
                               json=payload, timeout=10)
            if r.status_code == 200:
                self.root.after(0, lambda: self.status_label.config(text="Device registered."))
        except Exception:
            self.root.after(0, lambda: self.status_label.config(text="Server unavailable — offline mode."))

    def _start_heartbeat(self):
        def _loop():
            while True:
                try:
                    payload = {
                        "device_id": self.device_id,
                        "hostname":  platform.node(),
                        "os_info":   f"{platform.system()} {platform.release()}",
                        "user_info": {"app_version": "2.0.0"}
                    }
                    requests.post(f"{self.server_url}/api/device_heartbeat",
                                   json=payload, timeout=5)
                except Exception:
                    pass
                time.sleep(60)
        threading.Thread(target=_loop, daemon=True).start()

    def _start_license_monitoring(self):
        while True:
            try:
                r = requests.post(
                    f"{self.server_url}/api/device/validate_license",
                    json={"device_id":   self.device_id,
                          "license_key": self._get_local_license_key()},
                    timeout=10)
                if r.status_code == 200:
                    result = r.json()
                    if result.get("valid") and not self.license_data:
                        self.license_data = result
                        self.root.after(0, self._unlock_app)
                        return
            except Exception:
                pass
            time.sleep(30)

    def _unlock_app(self):
        self.license_status_label.stop(GREEN)
        self.license_status_label.config(
            text="✅  License Active — App Unlocked!", fg=GREEN)
        self.status_label.config(
            text="License approved! Loading main interface…", fg=GREEN)
        self.root.after(2000, self.create_main_ui)

    def _check_license(self):
        def _do():
            try:
                r = requests.post(
                    f"{self.server_url}/api/device/validate_license",
                    json={"device_id":   self.device_id,
                          "license_key": self._get_local_license_key()},
                    timeout=10)
                if r.status_code == 200:
                    result = r.json()
                    if result.get("valid"):
                        self.license_data = result
                        self.root.after(0, self._unlock_app)
                        return
            except Exception:
                pass
            local_key = self._get_local_license_key()
            if local_key:
                self.root.after(0, lambda: (
                    self.license_status_label.config(
                        text="⏳  License Pending — Waiting for approval…",
                        fg=ORANGE),
                    self.status_label.config(
                        text="Local license found — awaiting server confirmation.")))
            else:
                self.root.after(0, lambda: (
                    self.license_status_label.stop(RED),
                    self.license_status_label.config(
                        text="🔒  No License — Payment Required", fg=RED),
                    self.status_label.config(
                        text="Please complete payment to activate AriTyper.")))
        threading.Thread(target=_do, daemon=True).start()

    def _get_local_license_key(self):
        try:
            if os.path.exists("license.json"):
                with open("license.json") as f:
                    return json.load(f).get("license_key")
        except Exception:
            pass
        return None

    def submit_transaction(self):
        txn = self.transaction_entry.get().strip()
        if not txn or "Paste" in txn:
            self._set_txn_status("⚠  Please enter your transaction ID.", ORANGE)
            return
        if len(txn) < 4:
            self._set_txn_status("⚠  Transaction ID is too short.", ORANGE)
            return

        self.submit_btn.config(state="disabled", text="Submitting…")
        self._set_txn_status("Sending to server…", MUTED)

        def _thread():
            try:
                payload = {
                    "device_id":      self.device_id,
                    "transaction_id": txn,
                    "phone_number":   "66562536",
                    "amount":         "10000",
                    "network":        "mtn",
                    "notes":          f"License request from device {self.device_id}"
                }
                r = requests.post(f"{self.server_url}/api/submit_payment",
                                   json=payload, timeout=10)
                if r.status_code == 200:
                    result = r.json()
                    if result.get("success"):
                        self.root.after(0, lambda: self._set_txn_status(
                            "✅  Submitted! Awaiting approval…", GREEN))
                        self.root.after(0, lambda: messagebox.showinfo(
                            "Payment Submitted",
                            "Your payment has been submitted for approval.\n\n"
                            "The app will automatically unlock once approved.\n"
                            "This usually takes a few minutes."))
                        self.root.after(0, lambda: self.status_label.config(
                            text="Payment submitted — waiting for approval."))
                    else:
                        msg = result.get("message", "Unknown error")
                        self.root.after(0, lambda: (
                            self._set_txn_status(f"✗  {msg}", RED),
                            messagebox.showerror("Error", msg)))
                else:
                    self.root.after(0, lambda: (
                        self._set_txn_status("✗  Server error.", RED),
                        messagebox.showerror("Server Error", "Failed to submit payment.")))
            except Exception as e:
                self.root.after(0, lambda: (
                    self._set_txn_status(f"✗  Network error.", RED),
                    messagebox.showerror("Network Error", f"Failed to connect: {e}")))
            finally:
                self.root.after(0, lambda: self.submit_btn.config(
                    state="normal", text="📤   Submit Transaction ID"))
        threading.Thread(target=_thread, daemon=True).start()

    def _set_txn_status(self, msg, color):
        self.txn_status.config(text=msg, fg=color)

    def browse_file(self):
        path = filedialog.askopenfilename(
            title="Select Document",
            filetypes=[("PDF", "*.pdf"), ("Word", "*.docx"),
                       ("Text", "*.txt"), ("All", "*.*")])
        if path:
            self.selected_file = path
            fname = os.path.basename(path)
            self.file_label.config(text=fname, fg=TEXT)
            self.main_status_label.config(text=f"File loaded: {fname}", fg=GREEN)

    def start_typing(self):
        # Get text content - try document display first, then file
        text_content = ""
        
        if hasattr(self, 'doc_text'):
            try:
                self.doc_text.config(state="normal")
                text_content = self.doc_text.get("1.0", tk.END).strip()
                self.doc_text.config(state="disabled")
            except:
                pass
        
        # If no content in display, try loading from file
        if not text_content and hasattr(self, 'selected_file') and self.selected_file:
            try:
                file_ext = os.path.splitext(self.selected_file)[1].lower()
                if file_ext == '.txt':
                    try:
                        with open(self.selected_file, 'r', encoding='utf-8') as f:
                            text_content = f.read()
                    except UnicodeDecodeError:
                        with open(self.selected_file, 'r', encoding='latin-1') as f:
                            text_content = f.read()
                elif file_ext == '.docx' and DOCX_SUPPORT:
                    doc = Document(self.selected_file)
                    text_content = '\n'.join([para.text for para in doc.paragraphs])
                elif file_ext == '.pdf' and PDF_SUPPORT:
                    with open(self.selected_file, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text_content = ''
                        for page in pdf_reader.pages:
                            text_content += page.extract_text() + '\n'
            except:
                pass
        
        if not text_content:
            messagebox.showwarning("No Content", "Please load or paste text content first.")
            return
            
        self.is_typing = True
        self.start_btn.config(state="disabled")
        self.stop_btn.config(state="normal")
        self.main_status_label.config(text="Typing in progress…", fg=ACCENT)
        self.progress_bar.set_pct(0)

        def _run():
            try:
                # Bring target window to front if selected
                if hasattr(self, 'target_window') and self.target_window and WINDOWS_SUPPORT:
                    try:
                        win32gui.ShowWindow(self.target_window, win32con.SW_RESTORE)
                        win32gui.SetForegroundWindow(self.target_window)
                        time.sleep(1)  # Give window time to focus
                    except:
                        pass
                
                speed = self.speed_var.get()
                delay = 0.05 * max(0.05, 1.05 - speed / 100)
                total_chars = len(text_content)
                
                # Type the content
                if WINDOWS_SUPPORT:
                    # Use Windows API for typing
                    for i, char in enumerate(text_content):
                        if not self.is_typing:
                            break
                        
                        try:
                            if char == '\n':
                                win32api.keybd_event(win32con.VK_RETURN, 0, 0, 0)
                                win32api.keybd_event(win32con.VK_RETURN, 0, win32con.KEYEVENTF_KEYUP, 0)
                            elif char == '\t':
                                win32api.keybd_event(win32con.VK_TAB, 0, 0, 0)
                                win32api.keybd_event(win32con.VK_TAB, 0, win32con.KEYEVENTF_KEYUP, 0)
                            else:
                                # Convert character to key code
                                vk = win32api.VkKeyScan(char)
                                if vk != -1:
                                    win32api.keybd_event(vk & 0xFF, 0, 0, 0)
                                    win32api.keybd_event(vk & 0xFF, 0, win32con.KEYEVENTF_KEYUP, 0)
                                else:
                                    # Handle special characters
                                    if char.isupper():
                                        win32api.keybd_event(win32con.VK_SHIFT, 0, 0, 0)
                                        vk = win32api.VkKeyScan(char.lower())
                                        if vk != -1:
                                            win32api.keybd_event(vk & 0xFF, 0, 0, 0)
                                            win32api.keybd_event(vk & 0xFF, 0, win32con.KEYEVENTF_KEYUP, 0)
                                        win32api.keybd_event(win32con.VK_SHIFT, 0, win32con.KEYEVENTF_KEYUP, 0)
                        except:
                            # Skip problematic characters
                            pass
                        
                        # Update progress
                        progress = int((i + 1) / total_chars * 100)
                        self.root.after(0, self._set_progress, progress)
                        time.sleep(delay)
                else:
                    # Fallback - just simulate progress
                    for i in range(total_chars):
                        if not self.is_typing:
                            break
                        progress = int((i + 1) / total_chars * 100)
                        self.root.after(0, self._set_progress, progress)
                        time.sleep(delay)
                
                if self.is_typing:
                    self.root.after(0, self._done_typing)
                    
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("Typing Error", f"Failed to type: {str(e)}"))
                self.root.after(0, self.stop_typing)
                
        threading.Thread(target=_run, daemon=True).start()

    def _set_progress(self, v):
        self.progress_bar.set_pct(v)
        self._pct_lbl.config(text=f"{v} %")

    def _done_typing(self):
        self.is_typing = False
        self.start_btn.config(state="normal")
        self.stop_btn.config(state="disabled")
        self.main_status_label.config(text="✅  Typing complete!", fg=GREEN)

    def stop_typing(self):
        self.is_typing = False
        self.start_btn.config(state="normal")
        self.stop_btn.config(state="disabled")
        self.progress_bar.set_pct(0)
        self._pct_lbl.config(text="0 %")
        self.main_status_label.config(text="Stopped.", fg=MUTED)

    def _update_speed_label(self, val=None):
        v = int(self.speed_var.get())
        self.speed_label.config(text=f"{v} %")

    def _generate_device_id(self):
        parts = []
        try:    parts.append(socket.gethostname())
        except: pass
        parts.append(platform.system())
        parts.append(platform.machine())
        try:
            import uuid
            parts.append(str(uuid.getnode()))
        except: pass
        h = hashlib.sha256("|".join(parts).encode()).hexdigest()
        return f"ARI-{h[:16].upper()}"

    # ═══════════════════════════════════════════
    #  ENHANCED FUNCTIONALITY
    # ═══════════════════════════════════════════
    def select_target_window(self):
        """Allow user to select a window to type into"""
        if not WINDOWS_SUPPORT:
            messagebox.showerror("Error", "Window selection requires Windows OS with pywin32 installed.")
            return
            
        try:
            def enum_windows_callback(hwnd, windows):
                try:
                    if win32gui.IsWindowVisible(hwnd) and win32gui.GetWindowText(hwnd):
                        windows.append((hwnd, win32gui.GetWindowText(hwnd)))
                except:
                    pass
                return True

            windows = []
            win32gui.EnumWindows(enum_windows_callback, windows)
            
            if not windows:
                messagebox.showinfo("No Windows", "No accessible windows found.")
                return

            # Create window selection dialog
            dialog = tk.Toplevel(self.root)
            dialog.title("Select Target Window")
            dialog.geometry("500x400")
            dialog.configure(bg=BG)
            dialog.transient(self.root)
            dialog.grab_set()

            tk.Label(dialog, text="Select the window where you want to type:",
                    font=FT_BODY, bg=BG, fg=TEXT).pack(pady=10)

            # Listbox with scrollbar
            list_frame = tk.Frame(dialog, bg=BG)
            list_frame.pack(fill="both", expand=True, padx=20, pady=10)

            scrollbar = tk.Scrollbar(list_frame)
            scrollbar.pack(side="right", fill="y")

            window_listbox = tk.Listbox(list_frame, font=FT_BODY, bg=CARD2, fg=TEXT,
                                       yscrollcommand=scrollbar.set, relief="flat", bd=0)
            window_listbox.pack(side="left", fill="both", expand=True)
            scrollbar.config(command=window_listbox.yview)

            for hwnd, title in windows:
                window_listbox.insert(tk.END, f"{title[:60]}...")

            def select_window():
                selection = window_listbox.curselection()
                if selection:
                    hwnd, title = windows[selection[0]]
                    self.target_window = hwnd
                    if hasattr(self, 'window_info'):
                        self.window_info.config(text=f"Target: {title[:40]}...", fg=GREEN)
                    dialog.destroy()
                    messagebox.showinfo("Success", f"Selected: {title}")

            tk.Button(dialog, text="Select", font=FT_BTN, bg=GREEN, fg=TEXT,
                     command=select_window).pack(pady=10)
                     
        except Exception as e:
            messagebox.showerror("Error", f"Window selection failed: {str(e)}")

    def paste_text(self):
        """Paste text from clipboard to document display"""
        try:
            self.doc_text.config(state="normal")
            self.doc_text.delete("1.0", tk.END)
            
            # Get clipboard content
            try:
                import win32clipboard
                win32clipboard.OpenClipboard()
                data = win32clipboard.GetClipboardData()
                win32clipboard.CloseClipboard()
                self.doc_text.insert("1.0", data)
                self.main_status_label.config(text="✅ Text pasted successfully", fg=GREEN)
                if hasattr(self, 'doc_status'):
                    self.doc_status.config(text=f"Text pasted - {len(data)} characters", fg=GREEN)
            except:
                # Fallback to tkinter clipboard
                try:
                    data = self.root.clipboard_get()
                    self.doc_text.insert("1.0", data)
                    self.main_status_label.config(text="✅ Text pasted successfully", fg=GREEN)
                    if hasattr(self, 'doc_status'):
                        self.doc_status.config(text=f"Text pasted - {len(data)} characters", fg=GREEN)
                except:
                    messagebox.showinfo("Info", "No text found in clipboard")
                    
            self.doc_text.config(state="disabled")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to paste text: {str(e)}")

    def clear_document(self):
        """Clear document display"""
        self.doc_text.config(state="normal")
        self.doc_text.delete("1.0", tk.END)
        self.doc_text.config(state="disabled")
        self.main_status_label.config(text="Document cleared", fg=MUTED)
        if hasattr(self, 'doc_status'):
            self.doc_status.config(text="No content loaded - paste text or select a file", fg=MUTED)

    def load_document(self, file_path):
        """Load and display document content"""
        try:
            # Store file path for typing
            self.selected_file = file_path
            
            # Update file label
            fname = os.path.basename(file_path)
            self.file_label.config(text=fname[:30] + "...", fg=TEXT)
            
            # Load content into document display
            content = ""
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.txt':
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                except UnicodeDecodeError:
                    with open(file_path, 'r', encoding='latin-1') as f:
                        content = f.read()
                        
            elif file_ext == '.docx':
                if not DOCX_SUPPORT:
                    messagebox.showerror("Error", "DOCX support not available. Please install python-docx.")
                    return
                try:
                    doc = Document(file_path)
                    content = '\n'.join([para.text for para in doc.paragraphs])
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to load DOCX file: {str(e)}")
                    return
                    
            elif file_ext == '.pdf':
                if not PDF_SUPPORT:
                    messagebox.showerror("Error", "PDF support not available. Please install PyPDF2.")
                    return
                try:
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        content = ''
                        for page in pdf_reader.pages:
                            content += page.extract_text() + '\n'
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to load PDF file: {str(e)}")
                    return
            else:
                messagebox.showerror("Error", "Unsupported file format.")
                return
            
            # Display content in text widget
            if hasattr(self, 'doc_text'):
                try:
                    self.doc_text.config(state="normal")
                    self.doc_text.delete("1.0", tk.END)
                    self.doc_text.insert("1.0", content)
                    self.doc_text.config(state="disabled")
                except Exception as e:
                    print(f"Error updating text widget: {e}")
            
            self.main_status_label.config(text=f"✅ Loaded: {fname}", fg=GREEN)
            if hasattr(self, 'doc_status'):
                self.doc_status.config(text=f"Loaded {fname} - {len(content)} characters", fg=GREEN)
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load document: {str(e)}")

    def browse_file(self):
        """Enhanced file browser with document loading"""
        file_types = [
            ("All Supported", "*.txt;*.docx;*.pdf"),
            ("Text Files", "*.txt"),
            ("Word Documents", "*.docx"),
            ("PDF Files", "*.pdf"),
            ("All Files", "*.*")
        ]
        
        file_path = filedialog.askopenfilename(
            title="Select Document",
            filetypes=file_types
        )
        
        if file_path:
            self.load_document(file_path)


# ═══════════════════════════════════════════════════════
def main():
    try:
        print("Starting AriTyper Activated…")
        root = tk.Tk()
        AriTyperActivated(root)
        print("AriTyper ready!")
        root.mainloop()
    except KeyboardInterrupt:
        sys.exit(0)
    except Exception as e:
        import traceback
        traceback.print_exc()
        try:
            messagebox.showerror("Startup Error",
                f"An error occurred while starting AriTyper:\n\n{e}")
        except: pass
        sys.exit(1)

if __name__ == "__main__":
    main()
